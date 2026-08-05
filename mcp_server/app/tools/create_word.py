"""Tool：產生 Word（.docx）文件。

用結構化參數（章節/段落/表格/圖片），不是讓模型自己寫程式碼跑。參數用 Pydantic
model 宣告巢狀欄位，理由見 create_excel.py 開頭的說明（MCP 工具的 JSON schema
不會讀函式 docstring，只有 Pydantic model 的欄位才會出現在模型看到的 schema 裡）。
"""
import base64
import io

from docx import Document
from docx.shared import Inches
from mcp_types import CallToolResult, TextContent
from pydantic import BaseModel, Field


class TableSpec(BaseModel):
    headers: list[str] = Field(default_factory=list, description="表格標題列（第一列），可留空")
    rows: list[list[str]] = Field(default_factory=list, description="表格資料列，每列是一個字串陣列")


class SectionSpec(BaseModel):
    heading: str = Field("", description="這個章節的標題，留空則不加標題")
    level: int = Field(1, description="標題階層，1 最大（如同 H1），數字越大字級越小")
    paragraphs: list[str] = Field(default_factory=list, description="這個章節底下的段落文字，依序列出")
    table: TableSpec | None = Field(None, description="這個章節要放的表格（選填）")
    image_base64: str | None = Field(None, description="這個章節要放的圖片，base64 編碼內容（選填）")
    image_width_inches: float = Field(5.0, description="圖片寬度（英吋），預設 5")


def _build(title: str, sections: list[SectionSpec]) -> bytes:
    doc = Document()
    if title:
        doc.add_heading(title, level=0)
    for sec in sections:
        if sec.heading:
            doc.add_heading(sec.heading, level=sec.level)
        for para in sec.paragraphs:
            doc.add_paragraph(para)
        if sec.table:
            headers, rows = sec.table.headers, sec.table.rows
            cols = len(headers) or (len(rows[0]) if rows else 1)
            table = doc.add_table(rows=1 if headers else 0, cols=cols)
            try:
                table.style = "Light Grid Accent 1"
            except KeyError:
                pass  # 樣式不存在就用預設，不要因為排版問題整個失敗
            if headers:
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = str(h)
            for row in rows:
                cells = table.add_row().cells
                for i, v in enumerate(row):
                    cells[i].text = str(v)
        if sec.image_base64:
            doc.add_picture(io.BytesIO(base64.b64decode(sec.image_base64)),
                            width=Inches(sec.image_width_inches))
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def register(server) -> None:
    @server.tool(
        name="create_word",
        description="產生 Word（.docx）文件，支援標題階層、段落、表格、圖片。"
                    "何時使用：使用者要求產出報告、公文、會議紀錄等正式文件。",
    )
    def create_word(filename: str, title: str = "", sections: list[SectionSpec] | None = None) -> CallToolResult:
        name = filename if filename.endswith(".docx") else filename + ".docx"
        data = _build(title, sections or [])
        return CallToolResult(
            content=[TextContent(type="text", text=f"已產生 Word 文件：{name}")],
            structured_content={
                "filename": name,
                "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "data_base64": base64.b64encode(data).decode(),
            },
        )
